"""
DOCX counterpart to extractor.py.

Tries the same "Unified ID Convention" plain-text strategy first
(id_convention.py, shared with the PDF path -- format-agnostic, since it
operates on a flattened text blob either way), before falling back to the
heading-style-driven strategies below. Confirmed by direct inspection
that both this document's PDF and DOCX exports carry the identical
Grade-9/10 leaked-Markdown-marker defect, so the same defensive
tokenization handles both without a DOCX-specific fixup.

Word documents almost always carry proper paragraph styles ("Heading 1",
"Heading 2", ...), which is a much more reliable structural signal than
font-size guessing -- effectively the DOCX equivalent of a PDF's table of
contents.

Like the PDF path, extraction is scoped to the curriculum's "Unit(s),
Topics and CLOs" section: a Heading 1 matching SECTION_HEADING_RE (see
extractor.py), extracting only until the next Heading 1. Inside that scope,
each Heading 2 is a Unit. python-docx's `document.paragraphs` silently
skips table content entirely, so we walk the document body in its real
order (paragraphs and tables interleaved) via `_iter_block_items` and
classify each table by shape rather than by which heading precedes it:

    cloCode | description ("Learners will be able to...") | bloomLevel |
    mandatory | Key Concept / Topic | Evidence of Learning

is a CLO table (detected by header keywords, not exact column count/order,
so re-ordered or slightly renamed columns still work) -- each row becomes a
structured CLO entry, and rows are grouped by their "Key Concept / Topic"
cell into Topics. A plain 2-column table (e.g. "subjectCode | BIO",
"titleEn | Cell Biology", ...) is a unit metadata table and its rows are
merged into the unit's `metadata` dict. Any other table is flattened into
the current unit's text so nothing is silently dropped, even though we
don't specifically model it. Heading 3+ (e.g. "5.1.1 Topics and Curriculum
Learning Outcomes") is informational only -- the table-shape test above is
what actually finds the CLO data, so we don't need to detect that
sub-heading precisely.

If no "Unit(s), Topics and CLOs" heading is found at all, we fall back to
the older whole-document behavior (Heading 1 -> Unit, Heading 2 -> Topic),
so documents that don't follow the new format still parse the same as
before. And if a document has no heading styles whatsoever, we fall back
further to treating it as a single unit, same as the PDF path does.
"""

from __future__ import annotations

import io
from typing import Optional

import docx
from docx.document import Document as _DocxDocument
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from app.services.curriculum_parser import id_convention
from app.services.curriculum_parser.extractor import (
    SECTION_HEADING_RE,
    TOPICS_MARKER_RE,
    _classify_body_text,
    _normalize_ws,
    _rows_as_clo_rows,
    _rows_as_metadata,
)


def _iter_block_items(document: _DocxDocument):
    """Yield each Paragraph/Table in the document in true document order."""
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def _heading_level(style_name: str) -> Optional[int]:
    name = (style_name or "").lower()
    if name == "title" or name.startswith("heading 1"):
        return 1
    if name.startswith("heading 2"):
        return 2
    if name.startswith("heading"):
        # Heading 3+ folds into the current unit rather than growing the
        # tree deeper than Unit -> Topic; table shape (not heading depth)
        # is what actually locates the CLO data.
        return 3
    return None


def _tokenize(document: _DocxDocument) -> list[dict]:
    """Flattens the document into an ordered list of
    {"type": "heading", "level", "title"} / {"type": "text", "text"} /
    {"type": "table", "rows"} blocks."""
    blocks: list[dict] = []
    for block in _iter_block_items(document):
        if isinstance(block, Paragraph):
            level = _heading_level(getattr(block.style, "name", None) or "")
            text = block.text.strip()
            if not text:
                continue
            if level is not None:
                blocks.append({"type": "heading", "level": level, "title": text})
            else:
                blocks.append({"type": "text", "text": text})
        elif isinstance(block, Table):
            rows = [[c.text.strip() for c in r.cells] for r in block.rows]
            blocks.append({"type": "table", "rows": rows})
    return blocks


def _group_clos_into_topics(clo_rows: list[dict]) -> list[dict]:
    """Dict-based equivalent of extractor.py's _group_clos_into_topics --
    groups CLO rows into Topics using their "Topic" cell, falling back to
    "Key Concept / Topic" (older, combined-column format) if there's no
    distinct Topic column."""
    topics: dict[str, dict] = {}
    order: list[str] = []
    for clo in clo_rows:
        key = clo.get("topic") or clo.get("keyConcept") or "General"
        if key not in topics:
            order.append(key)
            topics[key] = {
                "sequenceOrder": len(order),
                "titleEn": key,
                "keyConcepts": [],
                "learningOutcomes": [],
                "clos": [],
                "rawText": "",
                "page": None,
            }
        topic = topics[key]
        key_concept = clo.get("keyConcept")
        if key_concept and key_concept not in topic["keyConcepts"]:
            topic["keyConcepts"].append(key_concept)
        if clo.get("description"):
            topic["learningOutcomes"].append(clo["description"])
        topic["clos"].append(clo)
    return [topics[k] for k in order]


def extract_structure(
    docx_bytes: bytes,
    subject_code: str,
    grade_level: int,
    academic_year: str,
) -> dict:
    document = docx.Document(io.BytesIO(docx_bytes))

    # Strategy 0: "Unified ID Convention" format -- see id_convention.py
    # and the module docstring above. Checked first since it's neither
    # table-driven nor heading-hierarchy-driven, so it isn't something
    # the strategies below would ever detect on their own.
    full_text = "\n".join(p.text for p in document.paragraphs if p.text.strip())
    if id_convention.looks_like_id_convention(full_text):
        return id_convention.extract_grade(full_text, grade_level, subject_code, academic_year)

    warnings: list[str] = []
    blocks = _tokenize(document)

    if not blocks:
        return _single_unit_fallback(document, subject_code, grade_level, academic_year, warnings)

    section_start = None
    for i, b in enumerate(blocks):
        if b["type"] == "heading" and b["level"] == 1 and SECTION_HEADING_RE.match(_normalize_ws(b["title"])):
            section_start = i
            break

    if section_start is None:
        warnings.append(
            'No "Unit(s), Topics and CLOs" heading was found; used the legacy heading-style '
            "extraction instead."
        )
        return _extract_legacy(document, blocks, subject_code, grade_level, academic_year, warnings)

    section_end = len(blocks)
    for j in range(section_start + 1, len(blocks)):
        b = blocks[j]
        if b["type"] == "heading" and b["level"] == 1:
            section_end = j
            break

    units = _build_units_topics_clos(blocks[section_start + 1 : section_end], warnings)
    if not units:
        warnings.append(
            'Found a "Unit(s), Topics and CLOs" heading but no Heading 2 unit sub-headings '
            "under it; nothing was extracted from that section."
        )

    return {
        "subjectCode": subject_code,
        "gradeLevel": grade_level,
        "academicYear": academic_year,
        "extractionStrategy": "docx_units_topics_clos",
        "pageCount": None,
        "units": units,
        "warnings": warnings,
    }


# ---------------------------------------------------------------------------
# "Unit(s), Topics and CLOs" section (new, table-driven)
# ---------------------------------------------------------------------------

def _build_units_topics_clos(scoped: list[dict], warnings: list[str]) -> list[dict]:
    units: list[dict] = []
    current: Optional[dict] = None
    unit_index = 0

    def finalize(state: Optional[dict]) -> Optional[dict]:
        if state is None:
            return None
        unit_meta = state["meta"]
        title_en = unit_meta.get("titleEn") or state["heading_title"]
        try:
            number = int(unit_meta["unitNumber"]) if unit_meta.get("unitNumber") else state["index"]
        except ValueError:
            number = state["index"]

        unit: dict = {"number": number, "titleEn": title_en, "page": None, "topics": []}
        if unit_meta:
            unit["metadata"] = unit_meta

        if state["clos"]:
            unit["topics"] = _group_clos_into_topics(state["clos"])
        else:
            body = "\n".join(state["lines"])
            concepts, outcomes, raw_leftover = _classify_body_text(body)
            unit["topics"] = [
                {
                    "sequenceOrder": 1,
                    "titleEn": title_en,
                    "keyConcepts": concepts,
                    "learningOutcomes": outcomes,
                    "clos": [],
                    "rawText": raw_leftover[:5000],
                    "page": None,
                }
            ]
            warnings.append(
                f'Unit "{title_en}" has no CLO table; kept its text as a single topic '
                "for manual review."
            )
        return unit

    for b in scoped:
        if b["type"] == "heading" and b["level"] == 2 and not TOPICS_MARKER_RE.search(b["title"]):
            finalized = finalize(current)
            if finalized:
                units.append(finalized)
            unit_index += 1
            current = {"heading_title": b["title"], "index": unit_index, "meta": {}, "clos": [], "lines": []}
            continue

        if current is None:
            continue  # content before the first unit heading inside the section -- discarded

        if b["type"] == "table":
            clo_rows = _rows_as_clo_rows(b["rows"])
            if clo_rows is not None:
                current["clos"].extend(clo_rows)
            else:
                kv = _rows_as_metadata(b["rows"])
                if kv:
                    current["meta"].update(kv)
                else:
                    current["lines"].extend(" | ".join(r) for r in b["rows"] if any(r))
        elif b["type"] == "text":
            current["lines"].append(b["text"])
        elif b["type"] == "heading":
            # Heading 3+ ("5.1.1 Topics and Curriculum Learning Outcomes") --
            # informational only, folded into leftover lines.
            current["lines"].append(b["title"])

    finalized = finalize(current)
    if finalized:
        units.append(finalized)

    return units


# ---------------------------------------------------------------------------
# Legacy fallback: Heading 1 -> Unit, Heading 2 -> Topic (pre-existing behavior)
# ---------------------------------------------------------------------------

def _extract_legacy(
    document: _DocxDocument,
    blocks: list[dict],
    subject_code: str,
    grade_level: int,
    academic_year: str,
    warnings: list[str],
) -> dict:
    sections: list[dict] = []
    for b in blocks:
        if b["type"] == "heading":
            if b["level"] <= 2:
                sections.append({"level": b["level"], "title": b["title"], "lines": [], "clos": []})
            elif sections:
                sections[-1]["lines"].append(b["title"])
            else:
                sections.append({"level": 2, "title": b["title"], "lines": [], "clos": []})
        elif b["type"] == "text":
            if sections:
                sections[-1]["lines"].append(b["text"])
        elif b["type"] == "table":
            clo_rows = _rows_as_clo_rows(b["rows"])
            if clo_rows is not None:
                if sections:
                    sections[-1]["clos"].extend(clo_rows)
            elif sections:
                sections[-1]["lines"].extend(" | ".join(r) for r in b["rows"] if any(r))

    if not sections:
        return _single_unit_fallback(document, subject_code, grade_level, academic_year, warnings)

    units: list[dict] = []
    unit_number = 0
    topic_seq = 0
    current_unit: Optional[dict] = None

    for sec in sections:
        body = "\n".join(sec["lines"])
        concepts, outcomes, raw_leftover = _classify_body_text(body)
        for clo in sec["clos"]:
            if clo["description"]:
                outcomes.append(clo["description"])
            if clo["keyConcept"] and clo["keyConcept"] not in concepts:
                concepts.append(clo["keyConcept"])

        has_content = bool(concepts or outcomes or raw_leftover or sec["clos"])

        if sec["level"] == 1 or current_unit is None:
            unit_number += 1
            topic_seq = 0
            current_unit = {"number": unit_number, "titleEn": sec["title"], "page": None, "topics": []}
            units.append(current_unit)
            if has_content:
                topic_seq += 1
                current_unit["topics"].append(
                    {
                        "sequenceOrder": topic_seq,
                        "titleEn": sec["title"],
                        "keyConcepts": concepts,
                        "learningOutcomes": outcomes,
                        "clos": sec["clos"],
                        "rawText": raw_leftover[:5000],
                        "page": None,
                    }
                )
        else:
            topic_seq += 1
            current_unit["topics"].append(
                {
                    "sequenceOrder": topic_seq,
                    "titleEn": sec["title"],
                    "keyConcepts": concepts,
                    "learningOutcomes": outcomes,
                    "clos": sec["clos"],
                    "rawText": raw_leftover[:5000],
                    "page": None,
                }
            )

    return {
        "subjectCode": subject_code,
        "gradeLevel": grade_level,
        "academicYear": academic_year,
        "extractionStrategy": "docx_heading_styles",
        "pageCount": None,
        "units": [u for u in units if u["topics"]],
        "warnings": warnings,
    }


def _single_unit_fallback(
    document: _DocxDocument,
    subject_code: str,
    grade_level: int,
    academic_year: str,
    warnings: list[str],
) -> dict:
    warnings.append(
        "No heading styles were found in the document; the whole "
        "document was kept as a single unit for manual review."
    )
    all_text = "\n".join(p.text for p in document.paragraphs if p.text.strip())
    concepts, outcomes, raw_leftover = _classify_body_text(all_text)
    return {
        "subjectCode": subject_code,
        "gradeLevel": grade_level,
        "academicYear": academic_year,
        "extractionStrategy": "docx_single_unit",
        "pageCount": None,
        "units": [
            {
                "number": 1,
                "titleEn": "Untitled document",
                "page": 0,
                "topics": [
                    {
                        "sequenceOrder": 1,
                        "titleEn": "Untitled document",
                        "keyConcepts": concepts,
                        "learningOutcomes": outcomes,
                        "clos": [],
                        "rawText": raw_leftover[:5000],
                        "page": 0,
                    }
                ],
            }
        ],
        "warnings": warnings,
    }
