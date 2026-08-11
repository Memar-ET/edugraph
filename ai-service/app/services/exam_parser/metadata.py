"""
Metadata-table extraction for the exam format that has a Subject/Grade
Level/Exam Type/Total Marks table at the top (the same "Extraction-Ready"
philosophy as curriculum's unit-metadata table, see
curriculum_parser/extractor.py's _rows_as_metadata). This is the
authoritative source for subject/grade/scope/marks when present -- more
reliable than guessing from the teacher-typed upload title, which is only
used as an upload-time fallback (see service/title_parser.go on the Go
side). service.py calls resolve_metadata() to correct the exam row once
this is extracted.
"""

from __future__ import annotations

import io
import re
from typing import Optional

import docx
import fitz  # PyMuPDF

GRADE_LEVEL_RE = re.compile(r"(\d{1,2})")
LEADING_INT_RE = re.compile(r"(\d+)")
FINAL_RE = re.compile(r"\bfinal\b", re.IGNORECASE)
MIDTERM_RE = re.compile(r"\bmid[- ]?term\b", re.IGNORECASE)
UNIT_TEST_RE = re.compile(r"\b(?:unit\s+test|unit\s+exam|quiz|test)\b", re.IGNORECASE)
# Mirrors backend/internal/assessment/service/title_parser.go's unitRangeRe --
# same pattern applied to the "Exam Type" cell instead of the title.
UNIT_RANGE_RE = re.compile(r"units?\s+(\d{1,2})\s*(?:[-–]\s*(\d{1,2}))?", re.IGNORECASE)

# Keys as they appear in the table, lowercased for matching.
KEY_ALIASES = {
    "subject": "subject",
    "grade level": "gradeLevel",
    "grade": "gradeLevel",
    "exam type": "examType",
    "total marks": "totalMarks",
}


def _rows_as_dict(rows: list[list[str]]) -> dict[str, str]:
    meta: dict[str, str] = {}
    for row in rows:
        cells = [c for c in ((c or "").strip() for c in row) if c]
        if len(cells) < 2:
            continue
        key = re.sub(r"\s+", " ", cells[0]).strip().lower()
        field = KEY_ALIASES.get(key)
        if field:
            meta[field] = re.sub(r"\s+", " ", " ".join(cells[1:])).strip()
    return meta


def extract_pdf_metadata_table(pdf_bytes: bytes, max_pages: int = 2) -> dict[str, str]:
    """Scans the first few pages for a table with Subject/Grade Level/Exam
    Type/Total Marks rows. Returns {} if none is found -- callers must treat
    that as "no correction available", not an error."""
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    try:
        for page_no in range(min(max_pages, doc.page_count)):
            try:
                finder = doc[page_no].find_tables()
            except Exception:
                continue
            for tab in finder.tables:
                rows = tab.extract()
                meta = _rows_as_dict(rows)
                if "subject" in meta and "gradeLevel" in meta:
                    return meta
        return {}
    finally:
        doc.close()


def extract_docx_metadata_table(docx_bytes: bytes) -> dict[str, str]:
    document = docx.Document(io.BytesIO(docx_bytes))
    for table in document.tables:
        rows = [[c.text for c in r.cells] for r in table.rows]
        meta = _rows_as_dict(rows)
        if "subject" in meta and "gradeLevel" in meta:
            return meta
    return {}


def derive_exam_scope(exam_type_text: str) -> Optional[str]:
    if FINAL_RE.search(exam_type_text):
        return "final_exam"
    if MIDTERM_RE.search(exam_type_text):
        return "midterm"
    if UNIT_TEST_RE.search(exam_type_text):
        return "unit_test"
    return None


def derive_unit_numbers(exam_type_text: str) -> Optional[list[int]]:
    """"Units 1-3" -> [1,2,3], "Unit 4" -> [4]. None (not []) when nothing
    matches -- distinguishes "not resolved" from "resolved, empty"."""
    m = UNIT_RANGE_RE.search(exam_type_text)
    if not m:
        return None
    start = int(m.group(1))
    end = int(m.group(2)) if m.group(2) else start
    if end < start:
        return [start]
    return list(range(start, end + 1))


def parse_metadata_table(raw: dict[str, str]) -> Optional[dict]:
    """Normalizes the raw {subject, gradeLevel, examType, totalMarks} table
    dict into {subjectHint, gradeLevel, examScope, totalMarks}. Returns None
    if any required field can't be resolved -- a partial correction would be
    worse than none."""
    if not raw:
        return None

    subject_hint = raw.get("subject", "").strip()
    if not subject_hint:
        return None

    grade_match = GRADE_LEVEL_RE.search(raw.get("gradeLevel", ""))
    if not grade_match:
        return None
    grade_level = int(grade_match.group(1))
    if not (1 <= grade_level <= 12):
        return None

    exam_scope = derive_exam_scope(raw.get("examType", ""))
    if exam_scope is None:
        return None

    marks_match = LEADING_INT_RE.search(raw.get("totalMarks", ""))
    if not marks_match:
        return None
    total_marks = int(marks_match.group(1))

    return {
        "subjectHint": subject_hint,
        "gradeLevel": grade_level,
        "examScope": exam_scope,
        "totalMarks": total_marks,
        "unitNumbers": derive_unit_numbers(raw.get("examType", "")),
    }
