"""
Answer Key table extraction (extends 2A, feeds 2C's MCQ auto-grading).

Teacher-authored "Extraction-Ready" exam documents end with an "Answer Key
& Marking Guide" table (Q# | CLO Code | Bloom | Marks | Correct Answer /
Marking Points) that extractor.py's STOP_MARKER_RE deliberately treats as
the end of the question list, not a question -- this module reads that
same table for its "Correct Answer" column instead. Student-facing copies
of an exam never have this table (by design, since it would reveal
answers), so extract_*_answer_key returning {} is the expected, common
case, not an error -- those questions simply stay ungraded until a teacher
grades them (same as essay/short-answer always does).

Only MCQ answers are extracted: a "Correct Answer" cell for a non-MCQ
question holds marking-guidance prose ("Two of: cell wall, ..."), not a
single option letter, and OPTION_LETTER_RE's word-boundary check naturally
rejects that (a leading word like "Discuss..." doesn't match "D\\b" since
there's no boundary between two word characters).
"""

from __future__ import annotations

import io
import re

import docx
import fitz  # PyMuPDF

QUESTION_NUMBER_RE = re.compile(r"^\s*(\d{1,3})\s*$")
OPTION_LETTER_RE = re.compile(r"^\s*([A-Da-d])\b")


def _find_column(header_cells: list[str], *hints: str) -> int | None:
    for idx, cell in enumerate(header_cells):
        lowered = (cell or "").strip().lower()
        for hint in hints:
            if hint in lowered:
                return idx
    return None


def _rows_to_answer_key(rows: list[list[str]]) -> dict[int, str]:
    if len(rows) < 2:
        return {}
    header = [c or "" for c in rows[0]]
    q_col = _find_column(header, "q#", "question")
    ans_col = _find_column(header, "correct answer", "answer")
    if q_col is None or ans_col is None:
        return {}

    out: dict[int, str] = {}
    for row in rows[1:]:
        cells = [(c or "").strip() for c in row]
        if q_col >= len(cells) or ans_col >= len(cells):
            continue
        q_match = QUESTION_NUMBER_RE.match(cells[q_col])
        letter_match = OPTION_LETTER_RE.match(cells[ans_col])
        if q_match and letter_match:
            out[int(q_match.group(1))] = letter_match.group(1).upper()
    return out


def extract_pdf_answer_key(pdf_bytes: bytes) -> dict[int, str]:
    """{sequenceNumber: correctOptionLetter}. Scans every table on every
    page and merges matches -- only a table whose header actually contains
    a Q#-like and Correct-Answer-like column contributes anything, so this
    can't cross-contaminate with the unrelated metadata table."""
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    try:
        out: dict[int, str] = {}
        for page in doc:
            try:
                finder = page.find_tables()
            except Exception:
                continue
            for tab in finder.tables:
                out.update(_rows_to_answer_key(tab.extract()))
        return out
    finally:
        doc.close()


def extract_docx_answer_key(docx_bytes: bytes) -> dict[int, str]:
    document = docx.Document(io.BytesIO(docx_bytes))
    out: dict[int, str] = {}
    for table in document.tables:
        rows = [[c.text for c in r.cells] for r in table.rows]
        out.update(_rows_to_answer_key(rows))
    return out
