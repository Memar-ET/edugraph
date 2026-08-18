"""Generate UI.docx — the complete EduGraph AI frontend specification.

Covers design system, all 6 roles, every page (built + planned), component
library, routing, accessibility, i18n, and implementation priorities.

Run: python scripts/gen_ui_spec.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── colour palette matching The Register design system ───────────────────
INK      = RGBColor(0x1D, 0x34, 0x50)   # primary-700
INK_MID  = RGBColor(0x28, 0x47, 0x69)   # primary-600
LEDGER   = RGBColor(0x4B, 0x44, 0x36)   # gray-700
SEAL     = RGBColor(0x96, 0x69, 0x2A)   # seal-600
HEALTH   = RGBColor(0x23, 0x5A, 0x38)   # health-600
ALERT    = RGBColor(0x83, 0x2F, 0x1F)   # alert-600
WHITE    = RGBColor(0xFF, 0xFF, 0xFF)
LGREY    = RGBColor(0xEF, 0xEB, 0xE1)   # gray-100

doc = Document()
sty = doc.styles["Normal"]
sty.font.name = "Calibri"
sty.font.size = Pt(10.5)

# ─── helpers ──────────────────────────────────────────────────────────────
def shade_cell(cell, hexcolor: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)

def h1(text, color=INK):
    p = doc.add_heading(text, level=1)
    for r in p.runs: r.font.color.rgb = color
    return p

def h2(text, color=INK_MID):
    p = doc.add_heading(text, level=2)
    for r in p.runs: r.font.color.rgb = color
    return p

def h3(text, color=LEDGER):
    p = doc.add_heading(text, level=3)
    for r in p.runs: r.font.color.rgb = color
    return p

def para(text, bold=False, italic=False, color=None, size=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold; r.italic = italic
    if color: r.font.color.rgb = color
    if size: r.font.size = Pt(size)
    return p

def bullet(label, body=""):
    p = doc.add_paragraph(style="List Bullet")
    if body:
        r = p.add_run(label); r.bold = True
        p.add_run(body)
    else:
        p.add_run(label)
    return p

def numbered(text):
    return doc.add_paragraph(text, style="List Number")

def tbl(headers, rows, widths=None, status_col=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hcells = t.rows[0].cells
    for i, h in enumerate(headers):
        hcells[i].text = ""
        r = hcells[i].paragraphs[0].add_run(h)
        r.bold = True; r.font.size = Pt(9); r.font.color.rgb = WHITE
        shade_cell(hcells[i], "1D3450")
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(val))
            r.font.size = Pt(9)
            if status_col is not None and i == status_col:
                v = str(val).lower()
                if any(x in v for x in ("built","done","exists","live","verified")):
                    r.font.color.rgb = HEALTH; r.bold = True
                elif any(x in v for x in ("partial","stub","unverified","wip")):
                    r.font.color.rgb = SEAL; r.bold = True
                elif any(x in v for x in ("missing","planned","todo","needed","broken")):
                    r.font.color.rgb = ALERT; r.bold = True
        if ri % 2 == 1:
            for c in cells: shade_cell(c, "F7F5EF")
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                if i < len(row.cells): row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t

def swatch(label, hexcolor, description):
    p = doc.add_paragraph()
    r = p.add_run("  ██  ")
    r.font.color.rgb = RGBColor(
        int(hexcolor[0:2],16), int(hexcolor[2:4],16), int(hexcolor[4:6],16))
    r.font.size = Pt(14)
    r2 = p.add_run(f"{label}  ")
    r2.bold = True; r2.font.size = Pt(10)
    r3 = p.add_run(f"#{hexcolor}  —  {description}")
    r3.font.size = Pt(9.5); r3.font.color.rgb = LEDGER

# ══════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title_p.add_run("EduGraph AI")
r.bold = True; r.font.size = Pt(36); r.font.color.rgb = INK

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub_p.add_run("Complete Frontend UI Specification")
r.font.size = Pt(18); r.font.color.rgb = INK_MID

meta_p = doc.add_paragraph()
meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta_p.add_run("Design System · Role-Based Access · All Pages · Component Library · Implementation Priorities\n")
r.font.size = Pt(11); r.font.color.rgb = LEDGER; r.italic = True
r2 = meta_p.add_run("2026-08-16  •  Curriculum-Intelligence Platform for Ethiopian K-12 Education")
r2.font.size = Pt(10); r2.font.color.rgb = LEDGER; r2.italic = True

doc.add_paragraph()
para(
    "This document is the authoritative frontend specification for EduGraph AI. "
    "It covers the design system, every role and the pages they can access, the "
    "component library, routing/auth-guard rules, accessibility requirements, "
    "internationalisation (English + Amharic), and a prioritised build list showing "
    "what is already live and what still needs implementation. "
    "It is intended for designers, frontend developers, and product reviewers.",
    italic=True, color=LEDGER,
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# SECTION 1 — DESIGN SYSTEM
# ══════════════════════════════════════════════════════════════════════════
h1("1.  Design System — \"The Register\"")
para(
    "The visual identity is named \"The Register\" — drawn from the world of "
    "official ministry records, curriculum registers, and exam ledger books. "
    "Everything references paper, ink, and the wax-seal stamp that makes a "
    "record official. This grounds the platform in the physical world Ethiopian "
    "teachers and administrators already trust, rather than importing a generic "
    "SaaS blue aesthetic.",
)
doc.add_paragraph()

h2("1.1  Palette")
para("Five semantic colour scales. Tailwind config overrides `primary` and `gray` "
     "so every existing utility class inherits the identity automatically.", italic=True, color=LEDGER)
doc.add_paragraph()

swatch("Ink (primary)", "1D3450", "Brand / interactive — deep desaturated blue, fountain-pen ink on a register page. Use for all interactive elements, links, active nav.")
swatch("Ledger (gray)", "4B4436", "Warm neutral — ledger paper. Replaces the default cool gray. Use for backgrounds (gray-50 #F7F5EF), body text (gray-900), borders.")
swatch("Seal (accent)", "96692A", "Ministry stamp. Used sparingly for approvals, certifications, confirmations, one warm highlight per view.")
swatch("Health (success/mastery)", "235A38", "Mastery, coverage, 'this is fine' data states. Gap closed, skill mastered, quality score good.")
swatch("Alert (warning/gap)", "832F1F", "Gaps, root causes, 'reteach now' data states. Never pure red — warm terracotta keeps it readable and non-alarming.")
doc.add_paragraph()

h2("1.2  Typography")
tbl(
    ["Scale", "Font", "Use"],
    [
        ["Display headings (h1/h2/h3)", "Fraunces (serif)", "Page titles, section headers, the EduGraph wordmark. Evokes printed registers and official documents."],
        ["Body + UI", "IBM Plex Sans + IBM Plex Sans Ethiopic", "All UI labels, paragraphs, table cells. Ethiopic variant ensures Amharic text renders on the same axis."],
        ["Mono / code", "IBM Plex Mono", "IDs, CLO codes, prerequisite edge data, JSON previews."],
    ],
    widths=[1.8, 2.2, 3.0],
)
para("Both IBM Plex variants are loaded from Google Fonts in the same @import rule so Amharic "
     "text in a mixed EN/AM paragraph never falls back to a system font with a different x-height.",
     italic=True, color=LEDGER)

h2("1.3  Spacing & Layout")
bullet("Base unit: ", "4 px (Tailwind default). All spacing in multiples of 4.")
bullet("Sidebar width: ", "256 px (w-64) fixed; collapses to an overlay on mobile.")
bullet("Page max-width: ", "None for dashboards (full bleed); 768 px for single-column forms (upload, login).")
bullet("Content padding: ", "px-6 py-6 on page containers; px-4 on sidebar.")
bullet("Card rounding: ", "rounded-xl (12 px) for cards, rounded-lg (8 px) for inputs, rounded-full for pills.")

h2("1.4  Elevation & Shadow")
bullet("stamp shadow: ", "`0 1px 2px rgb(15 24 38 / 0.06), 0 1px 1px rgb(15 24 38 / 0.04)` — the bespoke token. Applied to all cards, sidebar, header.")
bullet("No heavy drop-shadows. ", "This is a record system, not a consumer app. Elevation is conveyed by border + subtle stamp shadow, not large blur radii.")

h2("1.5  Motion")
bullet("Micro-interactions only: ", "nav active-rule slide (200 ms ease), sidebar mobile translate (200 ms), status-pill fade-in.")
bullet("prefers-reduced-motion: ", "All transitions and animations collapse to 0.01 ms. Already set in index.css.")
bullet("No page-transition animations. ", "Data density matters more than choreography on a government platform.")

h2("1.6  Dark Mode")
para("Not in the current build. The Tailwind config has no `darkMode` key. "
     "The ledger-paper warm neutrals make a future dark mode non-trivial — "
     "plan as a separate design track if required by MoE.", italic=True, color=SEAL)

h2("1.7  Accessibility Baseline")
bullet("WCAG 2.1 AA minimum. ", "All interactive colour pairs pass 4.5:1 contrast (primary-700 on gray-50 = 10.8:1).")
bullet("Keyboard navigation: ", "Full tab order, visible focus ring (ring-2 ring-primary-500 ring-offset-2) on all interactive elements.")
bullet("ARIA: ", "role, aria-label, aria-live regions on async loading states (exam submission, gap analysis polling).")
bullet("Language: ", "html lang=\"am\" or \"en\" switchable; both IBM Plex variants cover both scripts.")
bullet("Forms: ", "Every input has an associated <label> (htmlFor). Error messages use aria-describedby.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# SECTION 2 — COMPONENT LIBRARY
# ══════════════════════════════════════════════════════════════════════════
h1("2.  Component Library")
para("All components live in frontend/src/components/. "
     "Primitive UI atoms use Radix UI under the hood (no class-variance-authority yet — raw Tailwind). "
     "Charts use Recharts. The knowledge graph uses ReactFlow + dagre.")

h2("2.1  Primitive UI Components  (src/components/ui/)")
tbl(
    ["Component", "File", "Status", "Notes"],
    [
        ["Button", "button.tsx", "Built", "variant: primary | ghost | outline | destructive; size: sm | md | lg"],
        ["Input", "input.tsx", "Built", "Controlled, with label and error-message slot"],
        ["Label", "label.tsx", "Built", "Wraps Radix Label; htmlFor wired"],
        ["Select", "select.tsx", "Built", "Radix Select, styled to The Register palette"],
        ["Spinner", "spinner.tsx", "Built", "Inline loading indicator, respects reduced-motion"],
        ["Card", "card.tsx", "Built", "stamp shadow, rounded-xl, bg-white"],
        ["StatusPill", "status-pill.tsx", "Built", "Colour-coded: pending/parsing/parsed/approved/rejected etc."],
        ["EmptyState", "empty-state.tsx", "Built", "Centred illustration slot + heading + CTA"],
        ["Banner", "banner.tsx", "Built", "Info / warning / error stripe"],
        ["Seal", "seal.tsx", "Built", "Ministry approval stamp icon component"],
        ["ThreeDCard", "ThreeDCard.tsx", "Built", "Perspective-tilt card for hero/landing surfaces"],
        ["Modal / Dialog", "—", "Planned", "Radix Dialog, needed for confirm-close, reject-snapshot, misconception review"],
        ["Toast / Snackbar", "—", "Planned", "Sonner or Radix Toast for async feedback (save draft, submit exam)"],
        ["Tabs", "—", "Planned", "Radix Tabs for exam review (Questions | Quality | Insights)"],
        ["Combobox / Autocomplete", "—", "Planned", "Topic picker on Prerequisites page (already works with a Select; needs search for 500+ topics)"],
        ["DatePicker", "—", "Planned", "Exam scheduling (exam lifecycle publish → scheduled flow)"],
        ["FileDropzone", "—", "Planned", "Drag-and-drop upload on ExamUploadPage / UploadPage"],
        ["Tooltip", "—", "Planned", "Inline help text on quality scores, CLO codes, edge types"],
        ["Pagination", "—", "Planned", "Curriculum dashboard job list, exam list, student roster"],
    ],
    widths=[1.5, 1.5, 0.9, 3.0],
    status_col=2,
)

h2("2.2  Chart Components  (src/components/charts/)")
tbl(
    ["Component", "Status", "Used by"],
    [
        ["HeatmapGrid", "Built — onTopicClick wired to ExplainPage", "TeacherDashboardPage, SchoolAdminDashboardPage"],
        ["ScoreGauge", "Built", "SchoolAdminDashboardPage, QualityScoreGrid"],
        ["DistributionBars", "Built", "ExamQualityPage, MinistryDashboardPage"],
        ["DistributionDonutChart", "Built", "SchoolAdminDashboardPage"],
        ["PerformanceAreaChart", "Built", "TeacherDashboardPage, RegionalDashboardPage"],
        ["StatMetricCard", "Built", "All dashboards"],
        ["SparklineMini", "Planned", "Student subject-profile cards (7-day trend)"],
        ["RadarChart (CLO coverage)", "Planned", "ExamQualityPage — CLO axis vs mastery"],
        ["TreemapChart", "Planned", "MinistryDashboardPage — subject coverage treemap"],
        ["NetworkGraph (ReactFlow)", "Built — CurriculumGraphPage", "CurriculumGraphPage, PrerequisitesPage"],
    ],
    widths=[2.2, 2.2, 2.6],
    status_col=1,
)

h2("2.3  Dashboard Widgets  (src/components/dashboard/)")
tbl(
    ["Component", "Status", "Notes"],
    [
        ["StatMetricCard", "Built", "KPI tile: value, delta, icon, trend arrow"],
        ["ManagementTableCard", "Built", "Generic paginated table card"],
        ["PerformanceAreaChart", "Built", "Time-series line/area chart"],
        ["DistributionDonutChart", "Built", "Donut for categorical breakdowns"],
        ["ScheduleCalendarWidget", "Built", "Mini monthly calendar with exam event dots"],
        ["QualityScoreGrid", "Built (shared/)", "Grid of school quality score gauges"],
        ["NotificationBell", "Built (layout/)", "Polling /notifications with unread badge"],
    ],
    widths=[2.2, 1.2, 3.5],
    status_col=1,
)

h2("2.4  Layout Components  (src/components/layout/)")
tbl(
    ["Component", "Status", "Notes"],
    [
        ["AppShell", "Built", "Full sidebar + header + content slot. Sidebar nav grouped by section. Mobile overlay toggle."],
        ["AppHeader", "Built (legacy)", "Simple header; superseded by the sidebar brand header in AppShell. Can be deprecated."],
        ["NotificationBell", "Built", "Polls GET /notifications, shows unread count badge, dismisses on click"],
        ["nav-config.ts", "Built", "getNavItems(role) returns the correct nav items per role. Several items still point to '/' — see Section 4."],
    ],
    widths=[1.8, 1.0, 4.2],
    status_col=1,
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════
# SECTION 3 — AUTH & ROUTING
# ══════════════════════════════════════════════════════════════════════════
h1("3.  Authentication & Routing")

h2("3.1  Auth Flow")
bullet("JWT RS256 — ", "15-min access token + 7-day HttpOnly refresh cookie. Auth interceptor silently retries a 401 once before surfacing an error.")
bullet("Login landing — ", "curriculum_officer goes to /curriculum (CurriculumDashboardPage); every other role goes to / (DashboardRouter).")
bullet("Logout — ", "POST /auth/logout revokes the server-side session cookie, then clearAuth() wipes Zustand state, then navigate to /login.")

h2("3.2  Route Guards")
tbl(
    ["Guard function", "Roles allowed", "Applied to"],
    [
        ["requireAuth()", "Any authenticated user", "/career/paths, /students/:id/topics/:id/explain, /state-snapshots"],
        ["requireStudentAccess()", "student", "All /student/* and /students/me/* routes"],
        ["requireTeacherAccess()", "teacher, school_admin", "All /teacher/* routes"],
        ["requireCurriculumAccess()", "curriculum_officer, ministry_admin", "All /curriculum/*, /model-governance, /ministry/curriculum"],
    ],
    widths=[2.2, 1.8, 3.0],
)

h2("3.3  Full Route Table")
tbl(
    ["Path", "Component", "Status"],
    [
        ["/login", "LoginPage", "Built"],
        ["/", "DashboardRouter (role dispatch)", "Built"],
        ["/student/exams", "StudentExamListPage", "Built"],
        ["/student/exams/available", "ExamAvailabilityPage", "Built"],
        ["/student/exams/:id/pre", "PreExamPage", "Built"],
        ["/student/exams/:id", "TakeExamPage (autosave + draft restore)", "Built"],
        ["/student/tutor", "TutorPage", "Built"],
        ["/students/me/skill-states", "StudentSkillStatePage", "Built"],
        ["/students/:id/topics/:topicId/explain", "ExplainPage", "Built"],
        ["/student/study-plans", "StudyPlanPage", "Planned"],
        ["/student/subject-profiles", "SubjectProfilePage", "Planned"],
        ["/student/settings", "StudentSettingsPage", "Planned"],
        ["/teacher/exams/upload", "ExamUploadPage", "Built"],
        ["/teacher/exams/:id", "ExamReviewPage", "Built"],
        ["/teacher/exams/:id/grade", "GradeExamPage", "Built"],
        ["/teacher/exams/:id/quality", "ExamQualityPage", "Built"],
        ["/teacher/misconceptions", "MisconceptionReviewPage", "Built"],
        ["/teacher/exams", "TeacherExamListPage", "Planned"],
        ["/teacher/students", "TeacherStudentRosterPage", "Planned"],
        ["/teacher/students/:id", "TeacherStudentDetailPage", "Planned"],
        ["/teacher/class-analytics", "ClassAnalyticsPage", "Planned"],
        ["/school/quality", "SchoolQualityPage", "Planned"],
        ["/school/teachers", "TeacherRosterPage", "Planned"],
        ["/school/students", "StudentRosterPage", "Planned"],
        ["/school/reports", "SchoolReportsPage", "Planned"],
        ["/regional/schools", "RegionalSchoolsPage", "Planned"],
        ["/regional/analytics", "RegionalAnalyticsPage", "Planned"],
        ["/ministry/curriculum", "MinistryCurriculumPage", "Built"],
        ["/ministry/curriculum/:code", "MinistryCurriculumDetailPage", "Built"],
        ["/ministry/reports", "MinistryReportsPage", "Planned"],
        ["/ministry/regions", "RegionManagementPage", "Planned"],
        ["/curriculum", "CurriculumDashboardPage", "Built"],
        ["/curriculum/upload", "UploadPage", "Built"],
        ["/curriculum/jobs/:jobId", "JobReviewPage", "Built"],
        ["/curriculum/prerequisites", "PrerequisitesPage", "Built"],
        ["/curriculum/versions", "CurriculumVersionsPage", "Built"],
        ["/curriculum/subjects/:code/graph", "CurriculumGraphPage", "Built"],
        ["/model-governance", "ModelGovernancePage", "Built"],
        ["/curriculum/subjects/:code/qmatrix-quality", "QMatrixQualityPage", "Planned"],
        ["/career/paths", "CareerPathsPage", "Built"],
        ["/settings", "SettingsPage", "Planned"],
    ],
    widths=[2.6, 2.2, 0.9],
    status_col=2,
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# SECTION 4 — ROLE DEEP-DIVES
# ══════════════════════════════════════════════════════════════════════════
h1("4.  Role-Based Access — Feature Deep-Dives")

h2("4.1  Student")
para("Primary persona: K-12 learner preparing for, taking, and reviewing exams; "
     "receiving AI-generated personalised study support.", bold=True)
tbl(
    ["Feature", "Page", "Status", "API"],
    [
        ["Dashboard (subject health + study plan summary)", "StudentDashboardPage", "Built", "GET /students/me/subject-profiles"],
        ["Browse published exams", "ExamAvailabilityPage", "Built", "GET /exams"],
        ["Pre-exam briefing", "PreExamPage", "Built", "GET /exams/:id"],
        ["Take exam (autosave + draft restore + idempotency)", "TakeExamPage", "Built", "GET /exams/:id/questions, POST autosave/draft/submit"],
        ["Post-exam gap-analysis insight polling", "TakeExamPage (inline)", "Built", "GET /exams/:id/my-insight"],
        ["EG-GCKT knowledge map (skill states)", "StudentSkillStatePage", "Built", "GET /students/me/skill-states"],
        ["5-part topic explanation (EG-GCKT)", "ExplainPage", "Built", "GET /students/:id/topics/:topicId/explain"],
        ["Ask AI Tutor (Graph-RAG)", "TutorPage", "Built", "POST /tutor/ask"],
        ["Career path matches + generate", "CareerPathsPage", "Built", "GET + POST /students/me/career/*"],
        ["Active study plan (day-by-day blocks)", "StudyPlanPage", "Planned", "GET + POST /students/me/study-plans"],
        ["Subject health profiles (deep)", "SubjectProfilePage", "Planned", "GET /students/me/subject-profiles"],
        ["Account settings (language, password)", "SettingsPage", "Planned", "PATCH /auth/me"],
    ],
    widths=[2.1, 1.6, 0.9, 2.4],
    status_col=2,
)
bullet("Top missing: ", "StudyPlanPage, SubjectProfilePage, ExplainPage link-through from skill-state cards, SettingsPage.")

h2("4.2  Teacher")
para("Primary persona: classroom teacher managing exams, grading, reviewing AI quality "
     "and misconception alerts, monitoring class gaps.", bold=True)
tbl(
    ["Feature", "Page", "Status", "API"],
    [
        ["Class dashboard (heatmap + KPIs)", "TeacherDashboardPage", "Built", "GET /teachers/me/class-heatmap"],
        ["Heatmap topic click -> ExplainPage", "HeatmapGrid in dashboard", "Built", "GET /students/:id/topics/:topicId/explain"],
        ["Upload exam (PDF/DOCX)", "ExamUploadPage", "Built", "POST /exams/upload"],
        ["Review AI-parsed exam, edit, validate, publish, close", "ExamReviewPage", "Built", "Full exam lifecycle API"],
        ["Upload answer key", "ExamReviewPage", "Built", "POST /exams/:id/answer-key"],
        ["Bulk grade exam", "GradeExamPage", "Built", "POST /exams/:id/grades/bulk"],
        ["Per-student gap insights for an exam", "ExamReviewPage (insights tab)", "Built", "GET /exams/:id/insights"],
        ["Exam quality report", "ExamQualityPage", "Built", "GET /exams/:id/quality"],
        ["Print exam sheet + answer key", "ExamReviewPage (print actions)", "Built", "GET /exams/:id/print*"],
        ["Misconception review queue", "MisconceptionReviewPage", "Built", "GET/PATCH /misconceptions"],
        ["My exam list", "TeacherExamListPage", "Planned", "GET /exams"],
        ["Student roster for my class", "TeacherStudentRosterPage", "Planned", "GET /students"],
        ["Individual student detail + skill states", "TeacherStudentDetailPage", "Planned", "GET /students/:id + explain"],
        ["Class analytics (performance over time)", "ClassAnalyticsPage", "Planned", "GET /teachers/me/class-heatmap + GET /exams"],
    ],
    widths=[2.1, 1.7, 0.9, 2.3],
    status_col=2,
)
bullet("Top missing: ", "TeacherExamListPage (teachers have no list view of their own exams), student roster + detail, class analytics over time.")

doc.add_page_break()
h2("4.3  School Admin")
para("Primary persona: school principal with full teacher exam access plus "
     "school-wide quality, staff management, and student roster.", bold=True)
tbl(
    ["Feature", "Page", "Status", "API"],
    [
        ["School dashboard (quality score + heatmap)", "SchoolAdminDashboardPage", "Built", "GET /schools/:id/quality-scores"],
        ["All teacher exam capabilities (shared role)", "Same exam pages", "Built", "Same as teacher"],
        ["School quality score detail + breakdown", "SchoolQualityPage", "Planned", "GET /schools/:id/quality-scores"],
        ["Teacher roster (create, edit, deactivate)", "TeacherRosterPage", "Planned", "GET + POST + PATCH /teachers"],
        ["Student roster (create, edit, delete)", "StudentRosterPage", "Planned", "GET + POST + PATCH + DELETE /students"],
        ["Async reports (school_monthly)", "SchoolReportsPage", "Planned", "POST + GET /reports"],
        ["Notifications (flagged-for-review alerts)", "NotificationBell", "Built", "GET /notifications"],
    ],
    widths=[2.3, 1.7, 0.9, 2.1],
    status_col=2,
)
bullet("Top missing: ", "All management pages (staff roster, student roster, quality detail, reports page). Backend fully wired for all of them.")

h2("4.4  Regional Admin")
para("Primary persona: regional education bureau comparing school performance, "
     "flagging underperformers, creating schools and accounts.", bold=True)
tbl(
    ["Feature", "Page", "Status", "API"],
    [
        ["Regional dashboard (KPIs, underperforming schools)", "RegionalDashboardPage", "Built", "GET /ministry/regions/:id/stats + underperforming"],
        ["School comparison drill-down", "RegionalSchoolsPage", "Planned", "GET /schools + GET /schools/:id/quality-scores"],
        ["Regional analytics over time", "RegionalAnalyticsPage", "Planned", "GET /ministry/regions/:id/stats"],
        ["Create and manage schools", "RegionalSchoolsPage", "Planned", "POST + PATCH /schools"],
        ["Create staff accounts", "Within school pages", "Planned", "POST /teachers + /students"],
    ],
    widths=[2.3, 1.7, 0.9, 2.1],
    status_col=2,
)
bullet("Top missing: ", "RegionalSchoolsPage and RegionalAnalyticsPage. The entire management surface beyond the dashboard is absent.")

doc.add_page_break()
h2("4.5  Ministry Admin")
para("Primary persona: MoE official with full national authority — curriculum approval, "
     "model governance, national reports, region management.", bold=True)
tbl(
    ["Feature", "Page", "Status", "API"],
    [
        ["National overview dashboard", "MinistryDashboardPage", "Built", "GET /ministry/overview"],
        ["Curriculum browser (all subjects)", "MinistryCurriculumPage", "Built", "GET /curriculum/subjects"],
        ["Subject detail + version lineage", "MinistryCurriculumDetailPage", "Built", "GET /curriculum/subjects/:code/versions"],
        ["Upload + review + approve curriculum", "UploadPage, JobReviewPage", "Built", "Full curriculum pipeline"],
        ["Prerequisite graph + CLO resync", "PrerequisitesPage", "Built", "POST /curriculum/prerequisites/resync + /clos/resync"],
        ["Knowledge graph visualisation", "CurriculumGraphPage", "Built", "GET /curriculum/subjects/:code/graph"],
        ["Model governance (promote/reject snapshots)", "ModelGovernancePage", "Built", "GET/POST /model-snapshots"],
        ["Career paths management", "CareerPathsPage", "Built", "GET + POST /career/paths"],
        ["Q-matrix quality report", "QMatrixQualityPage", "Planned", "GET /curriculum/subjects/:code/qmatrix-quality"],
        ["Generate national reports (async)", "MinistryReportsPage", "Planned", "POST + GET /reports"],
        ["Region management (CRUD)", "RegionManagementPage", "Planned", "GET/POST/PATCH/DELETE /regions"],
        ["Underperforming schools + AI curriculum insights", "MinistryDashboardPage (partial)", "Partial", "GET /ministry/regions/:id/underperforming + POST /ministry/curriculum-insights"],
    ],
    widths=[2.3, 1.7, 0.9, 2.1],
    status_col=2,
)
bullet("Top missing: ", "MinistryReportsPage, RegionManagementPage, QMatrixQualityPage. Nav items for reports and regions still point to '/'.")

h2("4.6  Curriculum Officer")
para("Primary persona: MoE curriculum specialist uploading raw documents, reviewing "
     "AI-extracted structure, and maintaining the prerequisite/Q-matrix graph.", bold=True)
tbl(
    ["Feature", "Page", "Status", "API"],
    [
        ["Job list dashboard (own upload history)", "CurriculumDashboardPage", "Built", "GET /curriculum/jobs"],
        ["Upload curriculum PDF/DOCX", "UploadPage", "Built", "POST /curriculum/upload"],
        ["Review + edit + approve parsed structure", "JobReviewPage", "Built", "GET + POST /curriculum/jobs/:id"],
        ["View original uploaded file (proxied)", "JobReviewPage file viewer", "Built", "GET /storage/files/:jobId"],
        ["Knowledge graph visualisation per subject", "CurriculumGraphPage", "Built", "GET /curriculum/subjects/:code/graph"],
        ["Prerequisite editor (add, type, validate, history)", "PrerequisitesPage", "Built", "Full prerequisite API"],
        ["Version lineage + supersede", "CurriculumVersionsPage", "Built", "GET + POST version endpoints"],
        ["Q-matrix quality report", "QMatrixQualityPage", "Planned", "GET /curriculum/subjects/:code/qmatrix-quality"],
        ["Prerequisite quality report (orphaned topics)", "QMatrixQualityPage or PrerequisitesPage", "Planned", "GET /curriculum/subjects/:code/prerequisite-quality"],
        ["Model governance (candidate review)", "ModelGovernancePage", "Built", "GET /model-snapshots/candidates"],
        ["Parsing analytics (job success rate, timing)", "Planned page", "Planned", "GET /curriculum/jobs (aggregate)"],
    ],
    widths=[2.3, 1.7, 0.9, 2.1],
    status_col=2,
)
bullet("Top missing: ", "QMatrixQualityPage (two backend endpoints already exist), parsing analytics page, ExplainPage link-through from graph view.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# SECTION 5 — NAV GAP AUDIT
# ══════════════════════════════════════════════════════════════════════════
h1("5.  Nav-Config Gap Audit")
para("Every nav item that currently points to '/' is a stub that needs wiring. "
     "This table is the complete punch-list extracted from nav-config.ts.")
tbl(
    ["Role", "Nav label", "Current 'to'", "Target route", "Priority"],
    [
        ["student", "Analytics", "/", "/student/subject-profiles", "High"],
        ["student", "Help & Support", "/", "/help", "Low"],
        ["student", "Settings", "/", "/settings", "Medium"],
        ["teacher", "Class Analytics", "/", "/teacher/class-analytics", "High"],
        ["teacher", "Help & Support", "/", "/help", "Low"],
        ["teacher", "Settings", "/", "/settings", "Medium"],
        ["school_admin", "Quality Score", "/", "/school/quality", "High"],
        ["school_admin", "Help & Support", "/", "/help", "Low"],
        ["school_admin", "Settings", "/", "/settings", "Medium"],
        ["regional_admin", "Regional Analytics", "/", "/regional/analytics", "High"],
        ["regional_admin", "Help & Support", "/", "/help", "Low"],
        ["regional_admin", "Settings", "/", "/settings", "Medium"],
        ["ministry_admin", "National Reports", "/", "/ministry/reports", "High"],
        ["ministry_admin", "Help & Support", "/", "/help", "Low"],
        ["ministry_admin", "Settings", "/", "/settings", "Medium"],
        ["curriculum_officer", "Parsing Analytics", "/", "/curriculum/analytics", "Medium"],
        ["curriculum_officer", "Help & Support", "/", "/help", "Low"],
        ["curriculum_officer", "Settings", "/", "/settings", "Medium"],
    ],
    widths=[1.4, 1.6, 0.9, 1.9, 0.9],
    status_col=4,
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# SECTION 6 — I18N & ACCESSIBILITY
# ══════════════════════════════════════════════════════════════════════════
h1("6.  Internationalisation & Accessibility")

h2("6.1  Language Support")
para("The platform targets Ethiopian schools where the primary instruction language "
     "is Amharic (or regional Ethiosemitic/Cushitic languages), but all current UI "
     "strings are in English. The design system is ready for Amharic:")
bullet("Font: ", "IBM Plex Sans Ethiopic is loaded alongside IBM Plex Sans in the same font stack. Mixed EN/AM paragraphs render on the same baseline.")
bullet("Direction: ", "Amharic is left-to-right; no RTL layout work needed.")
bullet("i18n library needed: ", "react-i18next is the standard choice for React. Add an i18n/ directory with en.json and am.json message catalogs, wrap the app in I18nextProvider, and replace string literals with t('key') calls.")
bullet("Gap analysis / study plan text: ", "The ai-service already generates bilingual EN/AM narrative in gap_analysis/llm.py. The frontend just needs to display it; no translation layer required for AI-generated content.")
bullet("Exam content: ", "Question text and CLO descriptions can be in Amharic. The font renders them correctly; no other change needed.")

h2("6.2  Accessibility Checklist")
tbl(
    ["Requirement", "Current state", "Action needed"],
    [
        ["Colour contrast (4.5:1 AA minimum)", "primary-700 on gray-50 = 10.8:1. Alert colours designed for readability.", "Verify status-pill and badge colours with axe or contrast checker."],
        ["Keyboard navigation", "Tab order follows DOM order. Most interactive elements are native button/input.", "Add visible focus ring (ring-2 ring-primary-500) to any custom div-based interactive elements."],
        ["Screen reader labels", "Most pages use semantic HTML.", "Add aria-label to icon-only buttons (close X, notification bell). Add aria-live='polite' to async result areas (gap insight polling, autosave status)."],
        ["Form labels", "Input/Label components have htmlFor wired.", "Audit new planned forms (roster create, report generate) on build."],
        ["Motion", "prefers-reduced-motion applied globally in index.css.", "No action needed."],
        ["Language declaration", "html lang not yet set dynamically.", "Set lang='am' or lang='en' on <html> based on user language preference."],
        ["Error states", "Banner component exists.", "All async error paths should render a Banner with a descriptive message, not just a spinner that never resolves."],
    ],
    widths=[1.8, 2.2, 3.0],
    status_col=1,
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# SECTION 7 — IMPLEMENTATION PRIORITIES
# ══════════════════════════════════════════════════════════════════════════
h1("7.  Implementation Priorities")
para("Ordered by user impact vs. implementation effort. "
     "All items marked 'Planned' have fully wired backend APIs — these are frontend-only builds.", bold=True)

h2("P0 — Blocking gaps (high impact, backend already done)")
tbl(
    ["Item", "Effort", "Unblocks"],
    [
        ["TeacherExamListPage (/teacher/exams) — teachers have no list of their own exams", "S", "Core teacher workflow"],
        ["Wire all stub nav items to their target routes (see Section 5)", "S", "All roles — nav items that go to '/' look broken"],
        ["StudyPlanPage — student day-by-day study schedule", "M", "Student retention loop; backend GET+POST fully wired"],
        ["ExplainPage link-through — button from StudentSkillStatePage cards", "S", "Makes EG-GCKT explain discoverable"],
        ["SubjectProfilePage — student subject health deep-dive", "S", "Backend GET /students/me/subject-profiles exists"],
        ["TeacherStudentRosterPage + TeacherStudentDetailPage", "M", "Teacher can drill from heatmap to individual student"],
    ],
    widths=[3.8, 0.6, 2.6],
)

h2("P1 — High value, straightforward builds")
tbl(
    ["Item", "Effort", "Notes"],
    [
        ["SchoolQualityPage — quality score breakdown for school_admin", "M", "Rich data already in /schools/:id/quality-scores"],
        ["StudentRosterPage + TeacherRosterPage for school_admin", "M", "Full CRUD API wired"],
        ["MinistryReportsPage — trigger + poll async reports", "M", "POST /reports/generate + poll GET /reports/:id (same pattern as gap insight polling)"],
        ["RegionManagementPage for ministry_admin", "M", "GET/POST/PATCH/DELETE /regions all wired"],
        ["QMatrixQualityPage — two tabs: Q-matrix quality + prerequisite quality", "M", "Two backend endpoints already exist"],
        ["ClassAnalyticsPage — teacher performance trend over time", "L", "Needs time-series from multiple exam results"],
        ["Modal / Dialog primitives", "S", "Needed for confirm-close exam, reject snapshot, misconception review forms"],
        ["Toast / Snackbar system (Sonner)", "S", "Autosave feedback exists in TakeExamPage; rest of app needs the same pattern"],
    ],
    widths=[2.8, 0.6, 3.6],
)

h2("P2 — Medium term")
tbl(
    ["Item", "Effort", "Notes"],
    [
        ["RegionalSchoolsPage + RegionalAnalyticsPage", "L", "Regional admin has essentially one page today"],
        ["SchoolReportsPage (school_admin async report trigger)", "S", "Same pattern as MinistryReportsPage"],
        ["SettingsPage (language toggle EN/AM, password change)", "S", "All roles share this page"],
        ["react-i18next setup + am.json catalog", "L", "Foundation for full Amharic UI"],
        ["FileDropzone component for upload pages", "S", "UX improvement on UploadPage and ExamUploadPage"],
        ["Combobox/autocomplete for topic picker", "M", "PrerequisitesPage topic picker breaks at 500+ topics with a plain Select"],
        ["DatePicker for exam scheduling", "S", "Needed when scheduled exam lifecycle state is surfaced"],
        ["Pagination component", "S", "Curriculum dashboard, student exam list, staff roster all need it"],
        ["SparklineMini chart for student subject cards", "S", "7-day mastery trend on StudentDashboardPage"],
        ["Tooltip component", "S", "CLO codes, edge types, quality score factors all need inline help text"],
        ["CurriculumGraphPage legend + filter controls", "M", "Graph is implemented but needs UX polish for large subjects"],
    ],
    widths=[2.8, 0.6, 3.6],
)

h2("P3 — Nice-to-have / post-launch")
tbl(
    ["Item", "Effort", "Notes"],
    [
        ["Dark mode", "XL", "Non-trivial given warm ledger palette — separate design track"],
        ["Offline mode UI (School Box)", "XL", "sync-agent is built; frontend needs an offline indicator + queue status"],
        ["Help & Support page", "M", "Static FAQ/documentation; low product urgency"],
        ["Print stylesheet for reports", "M", "National/school reports that print cleanly"],
        ["PWA / installable app", "L", "Useful for schools with limited connectivity; needs service worker"],
        ["RadarChart for CLO coverage on ExamQualityPage", "M", "Visual polish on an already-working page"],
    ],
    widths=[2.8, 0.6, 3.6],
)

h2("7.1  Missing Component Summary")
tbl(
    ["Component", "Needed by", "Priority"],
    [
        ["Modal / Dialog (Radix Dialog)", "Exam close confirm, reject snapshot, misconception forms", "P1"],
        ["Toast / Snackbar (Sonner)", "Autosave feedback, plan generation, career generate", "P1"],
        ["Tabs (Radix Tabs)", "ExamReviewPage (Questions | Quality | Insights), QMatrixQualityPage", "P1"],
        ["Combobox / Autocomplete", "PrerequisitesPage topic picker (500+ topics)", "P1"],
        ["FileDropzone", "UploadPage, ExamUploadPage", "P2"],
        ["DatePicker", "Exam scheduling", "P2"],
        ["Pagination", "Job list, exam list, rosters", "P2"],
        ["Tooltip", "CLO codes, edge types, quality factors", "P2"],
        ["SparklineMini", "Student dashboard subject cards", "P2"],
    ],
    widths=[1.8, 3.0, 0.8],
    status_col=2,
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# SECTION 8 — PAGE SPECS FOR TOP-PRIORITY PLANNED PAGES
# ══════════════════════════════════════════════════════════════════════════
h1("8.  Page Specs — Top-Priority Planned Pages")
para("Each spec covers the route, guard, layout, key data sources, and main UI elements. "
     "These are build-ready; the APIs are confirmed live.", italic=True, color=LEDGER)

h2("8.1  StudyPlanPage  (/student/study-plans)")
bullet("Guard: ", "requireStudentAccess")
bullet("Data: ", "GET /students/me/study-plans (active plan), POST /students/me/study-plans (generate new)")
bullet("Layout: ", "AppShell. If no active plan: EmptyState with a 'Generate my study plan' button (calls POST, then polls until status=ready). If plan exists: accordion of days (Day 1, Day 2, ...) each containing topic cards.")
bullet("Topic card: ", "Topic title, estimated hours, action type badge (practice/diagnostic/prereq-review from EG-GCKT action ranking), mastery status indicator (at_risk/learning/mastered), 'Explain this topic' link to ExplainPage.")
bullet("Header actions: ", "Regenerate plan button (confirm dialog: this will replace the current plan).")
bullet("Empty prereq state: ", "If the student has no gap records yet, show a prompt to take an exam first.")

h2("8.2  TeacherExamListPage  (/teacher/exams)")
bullet("Guard: ", "requireTeacherAccess")
bullet("Data: ", "GET /exams (filtered to caller's school/grade)")
bullet("Layout: ", "AppShell. Stat bar (total exams, drafts, published, closed). Table with columns: Exam name, Subject/Grade, Status (StatusPill), Questions, Submissions, Created date, Actions.")
bullet("Actions per row: ", "View (->ExamReviewPage), Grade (->GradeExamPage), Quality (->ExamQualityPage), Close (confirm dialog).")
bullet("Empty state: ", "EmptyState with an 'Upload exam' CTA.")
bullet("Upload button: ", "Top-right, navigates to /teacher/exams/upload.")

h2("8.3  TeacherStudentRosterPage  (/teacher/students)")
bullet("Guard: ", "requireTeacherAccess")
bullet("Data: ", "GET /students (server-side scoped to caller's school)")
bullet("Layout: ", "AppShell. Search input + grade filter. Table: name, grade, last active, mastery trend (SparklineMini), actions.")
bullet("Row click: ", "Navigates to /teacher/students/:id (TeacherStudentDetailPage).")

h2("8.4  TeacherStudentDetailPage  (/teacher/students/:id)")
bullet("Guard: ", "requireTeacherAccess")
bullet("Data: ", "GET /students/:id, GET /students/:id/topics/:topicId/explain (lazy per-topic)")
bullet("Layout: ", "AppShell. Student profile header (name, grade, school). Tab strip: Skill States | Exam History | Misconceptions.")
bullet("Skill States tab: ", "Same card grid as StudentSkillStatePage but for the selected student. 'Explain' button on each card.")
bullet("Exam History tab: ", "List of this student's submissions with scores.")
bullet("Misconceptions tab: ", "Any confirmed misconception_hypotheses for this student.")

h2("8.5  MinistryReportsPage  (/ministry/reports)")
bullet("Guard: ", "requireCurriculumAccess (ministry_admin)")
bullet("Data: ", "POST /reports/generate (trigger), GET /reports/:id (poll status)")
bullet("Layout: ", "AppShell. Two-column: left sidebar selects report type (school_monthly / national_heatmap / clo_coverage), right panel shows params form + history list.")
bullet("Params form: ", "Varies by report type. school_monthly: school picker + month selector. national_heatmap: no params. clo_coverage: subject code.")
bullet("Status poll: ", "Same pattern as gap-analysis insight polling in TakeExamPage — poll every 5s until status=ready, then show a download/view button.")
bullet("History list: ", "Table of past reports (type, params, status, generated_at, requester).")

h2("8.6  QMatrixQualityPage  (/curriculum/subjects/:code/qmatrix-quality)")
bullet("Guard: ", "requireCurriculumAccess")
bullet("Data: ", "GET /curriculum/subjects/:code/qmatrix-quality, GET /curriculum/subjects/:code/prerequisite-quality")
bullet("Layout: ", "AppShell. Subject code in page header. Two tabs: Q-Matrix Issues | Prerequisite Issues.")
bullet("Q-Matrix Issues tab: ", "Table of questions with missing / low-confidence / ambiguous skill mappings. Each row has an 'Edit mappings' link -> /questions/:id/skill-mappings (a future page).")
bullet("Prerequisite Issues tab: ", "Table of orphaned topics (no edges) and low-confidence edges. 'View in prerequisites editor' link per row.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# SECTION 9 — STATE MANAGEMENT
# ══════════════════════════════════════════════════════════════════════════
h1("9.  State Management")
tbl(
    ["Store", "File", "What it holds"],
    [
        ["auth.store", "src/stores/auth.store.ts", "user, isAuthenticated, tokens. clearAuth(), landingPathFor(), role guards. JWT interceptor reads from here."],
        ["ui.store", "src/stores/ui.store.ts", "Global UI flags (sidebar open, active theme, language preference)."],
        ["offline.store", "src/stores/offline.store.ts", "Offline/online status, pending sync queue length (School Box indicator)."],
        ["TanStack Query", "src/lib/query/", "All server state (exams, curriculum, skill states, notifications). Cached, deduplicated, background-refreshed. Do not duplicate in Zustand."],
    ],
    widths=[1.4, 2.2, 3.4],
)
bullet("Rule: ", "Server state lives in TanStack Query. Only global client-only state lives in Zustand. No component-local state for data that other components need.")
bullet("Query keys: ", "Defined in src/lib/query/keys.ts. Always use the factory pattern: examKeys.detail(id), not hardcoded strings.")
bullet("Mutations: ", "Use useMutation from TanStack Query. On success, invalidate the relevant query keys so the UI stays in sync without manual refetch calls.")

h1("10.  Summary")
para("The EduGraph AI frontend is substantially built — not a scaffold. "
     "The design system is complete (The Register palette, IBM Plex Sans + Ethiopic, warm ledger neutrals), "
     "the component library is solid, the AppShell is production-quality, and 29 pages "
     "across all 6 roles are implemented and type-checked.", bold=False)
doc.add_paragraph()
para("The real gap is that roughly 20 additional pages are needed to make every nav item "
     "functional and to serve the management workflows that school admins, regional admins, "
     "and ministry admins need. All of these pages have fully wired backend APIs. "
     "They are frontend-only builds — no new backend work is required.", italic=True, color=HEALTH)
doc.add_paragraph()
tbl(
    ["Category", "Built", "Planned"],
    [
        ["Pages (feature components)", "29", "~20"],
        ["UI primitive components", "11", "9"],
        ["Chart components", "6", "3"],
        ["Routes in router.tsx", "24", "~18"],
        ["Nav items wired to real routes", "~20", "~17 point to '/'"],
    ],
    widths=[2.5, 1.5, 1.5],
)
doc.add_paragraph()
para("End of UI specification.", italic=True, color=LEDGER)

doc.save(r"D:\EDUGRAPH PROJECT\edugraph\UI.docx")
print("UI.docx written OK")
