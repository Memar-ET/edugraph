"""Generate progress.docx — comprehensive EduGraph AI project status report.

Regenerated 2026-08-16 after the Production Completion Roadmap (Phases 0-12).
Run: python scripts/gen_progress.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1F, 0x3A, 0x5F)
BLUE = RGBColor(0x2E, 0x6D, 0xA4)
GREEN = RGBColor(0x2E, 0x7D, 0x32)
AMBER = RGBColor(0xB7, 0x6E, 0x00)
RED = RGBColor(0xC6, 0x28, 0x28)
GREY = RGBColor(0x55, 0x55, 0x55)

doc = Document()

# base style
base = doc.styles["Normal"]
base.font.name = "Calibri"
base.font.size = Pt(10.5)


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def h1(text, color=NAVY):
    p = doc.add_heading(text, level=1)
    for r in p.runs:
        r.font.color.rgb = color
    return p


def h2(text, color=BLUE):
    p = doc.add_heading(text, level=2)
    for r in p.runs:
        r.font.color.rgb = color
    return p


def para(text, bold=False, color=None, size=None, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color
    if size:
        r.font.size = Pt(size)
    return p


def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def table(headers, rows, widths=None, status_col=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(htext)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9.5)
        shade(hdr[i], "1F3A5F")
    for ridx, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(9)
            if status_col is not None and i == status_col:
                v = str(val).lower()
                if "done" in v or "built" in v or "pass" in v or "live" in v:
                    run.font.color.rgb = GREEN
                elif "partial" in v or "stub" in v or "unverified" in v:
                    run.font.color.rgb = AMBER
                elif "broken" in v or "missing" in v or "not" in v:
                    run.font.color.rgb = RED
                run.bold = True
        if ridx % 2 == 1:
            for c in cells:
                shade(c, "EEF3F9")
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t


# ============================== COVER ==============================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("EduGraph AI")
r.bold = True
r.font.size = Pt(34)
r.font.color.rgb = NAVY

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Project Progress & Production-Readiness Report")
r.font.size = Pt(16)
r.font.color.rgb = BLUE

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run("Curriculum-Intelligence Platform for Ethiopian K-12 Education\n")
r.font.size = Pt(11)
r.font.color.rgb = GREY
r = meta.add_run("Generated 2026-08-16  •  Branch: feature/eg-gckt-qa-validation")
r.font.size = Pt(10)
r.font.color.rgb = GREY
r.italic = True

doc.add_paragraph()
para(
    "This report captures the full arc of the project: the product idea, the system "
    "architecture, every implemented feature, what has been verified versus what remains "
    "unproven, the known gaps, and the roadmap still ahead. It supersedes the initial "
    "progress report and now reflects the Production Completion Roadmap (Phases 0-12) "
    "landed on this branch.",
    italic=True, color=GREY,
)
doc.add_page_break()

# ===================== 1. THE IDEA =====================
h1("1.  The Idea")
para(
    "EduGraph AI turns raw Ethiopian Ministry of Education curriculum documents into a "
    "living knowledge graph, then uses that graph to diagnose why individual students "
    "struggle, generate personalized study plans, and give teachers and administrators a "
    "curriculum-aware view of learning across schools, regions, and the nation."
)
para("The platform is built around a few core beliefs:", bold=True)
bullet("Curriculum is a graph, not a list. ", "Topics have prerequisites; a student fails a topic because an upstream prerequisite was never mastered. Modeling that graph is what makes root-cause diagnosis possible.")
bullet("Say 'unknown' rather than manufacture confidence. ", "The learner-modeling engines report insufficient_data honestly instead of fabricating benchmark numbers when real interaction data does not yet exist.")
bullet("Human-in-the-loop over full automation. ", "AI proposes (parsed structure, inferred prerequisites, misconception hypotheses); a curriculum officer or teacher confirms before anything becomes authoritative.")
bullet("Runs locally first. ", "The entire stack runs on a laptop via Docker Compose — a deliberate design choice so schools with limited connectivity (the 'School Box' deployment) are first-class, not an afterthought.")

# ===================== 2. ARCHITECTURE =====================
h1("2.  System Architecture")
para("EduGraph AI is a service-oriented stack orchestrated by Docker Compose. Postgres moved to hosted Supabase on 2026-08-10; everything else runs locally.")
table(
    ["Service", "Tech", "Role", "Port"],
    [
        ["api", "Go 1.22 (chi, manual DI)", "Monolith REST API, RBAC, storage adapter, Neo4j sync", "8080"],
        ["ai-service", "Python 3.12 FastAPI", "8 background workers + tutor HTTP route", "8000"],
        ["frontend", "React 18 + TS + Vite", "Role dashboards for all 6 roles", "5173"],
        ["postgres", "Supabase (hosted)", "System of record; local fallback via profile", "5432"],
        ["neo4j", "Neo4j (local)", "Curriculum + student-progress knowledge graph", "7687"],
        ["redis", "Redis (local)", "6 job queues + refresh tokens + quality cache", "6379"],
        ["metrics", "Go internal server", "JSON /metrics + /health, queue depth poller", "9090"],
        ["sync-agent", "Go binary", "School Box offline sync (outbox push / cloud pull)", "8081"],
    ],
    widths=[1.2, 1.9, 3.0, 0.6],
)
para("Backend layering", bold=True, color=BLUE)
para("internal/<domain>/{handler,service,repository,dto} per domain (auth, curriculum, assessment, career, student, teacher, school, region, ministry, sync, notification, jobs, storage, reports, modeling). pkg/ holds cross-cutting infra: crypto, storage, middleware (security, ratelimit, audit), metrics, telemetry, database clients, config, errors, validator, and the Go->ai-service HTTP client.")
para("AI service", bold=True, color=BLUE)
para("Not Celery. Seven Redis BRPOP consumers plus one time-based nightly worker (refit_worker), all started as FastAPI lifespan tasks. Six queues: curriculum:parse, exam:parse, exam:answerkey, gap:analyze, studyplan:generate, gckt:trace, plus report:generate (added this roadmap).")

# ===================== 3. THE CURRICULUM PIPELINE =====================
h1("3.  Core Feature — The Curriculum Pipeline")
para("A 4-step pipeline built on a dual-storage adapter (Postgres BYTEA in dev, S3 in prod — the S3 provider now exists) so the same handler/service code runs both locally and in the cloud with zero changes.")
table(
    ["Step", "What happens", "Where"],
    [
        ["1. Upload", "Role-gated, magic-byte validated, saved via StorageProvider, job pushed to Redis", "curriculum handler/service"],
        ["2. Parse", "Worker BRPOPs, extracts structure (PyMuPDF; TOC-first / font-heuristic / ID-convention Strategy 0)", "curriculum_worker"],
        ["3. Review", "Officer edits the unit/topic/CLO tree in an editable UI, views original file via blob-fetch", "JobReviewPage.tsx"],
        ["4. Finalize", "One transaction upserts subjects/units/topics/CLOs, mirrors Subject->Unit->Topic->CLO into Neo4j", "ApproveAndPromote"],
    ],
    widths=[0.9, 4.4, 1.4],
)
para("Real content ingested: the full Ethiopian MoE Biology G7-G12 curriculum — 6 subjects, 33 units, 554 topics (109 top-level + 445 subtopics), 806 CLOs, 4,593 topic-CLO mappings, and 90 cross-grade prerequisite edges — verified count-for-count across Postgres and Neo4j.", bold=True, color=GREEN)
doc.add_page_break()

# ===================== 4. IMPLEMENTED FEATURES =====================
h1("4.  Implemented Features")

h2("4.1  Assessment & Intelligence Pipeline (Phases 2-4)")
bullet("Exam upload / parse / grade ", "— teacher uploads, worker extracts questions, separate answer-key flow aligns keys to CLOs, students take exams, submissions graded and CLO-mastery-tracked (per-answer timeSpentSecs).")
bullet("Prerequisite graph write path ", "— recursive-CTE cycle rejection, validated-vs-inferred provenance, Neo4j mirror + bulk resync backstop.")
bullet("Gap Analysis Engine ", "— 3-pass pipeline (triage -> prerequisite root-cause walk -> one Gemini call), bilingual EN/AM.")
bullet("Study Plan Generator ", "— Kahn topological sort over prerequisite edges + root-cause-before-symptom ordering, greedy day-packing.")
bullet("Graph-RAG Tutor ", "— the only ai-service HTTP route; injects the student's own gap records + prerequisite chain into the prompt.")
bullet("Class-Wide Gap Heatmap ", "— Neo4j STRUGGLED_WITH mirror, cross-grade alert when >40% of a class struggles with a topic.")
bullet("Exam Quality + School Quality Score ", "— discrimination index, difficulty calibration, coverage; nightly recompute via Go goroutine.")

h2("4.2  EG-GCKT — Graph-Cognitive Knowledge Tracing")
para("A 12-milestone research-grade learner-modeling layer. Compiles/imports/type-checks; not yet exercised against a live stack (see gaps).", italic=True, color=AMBER)
bullet("Analytical engines ", "— real BKT (4-param Bayesian), DINA (joint posterior), 2PL IRT (Newton-Raphson MLE), all on honest population defaults, all scope-aware (check for an active model snapshot before falling back).")
bullet("Graph-Cognitive State Fusion ", "— combines evidence weighted by reliability x recency x sample-size; disagreement widens uncertainty; graph-reasoning is deliberately the lowest-weighted source.")
bullet("Root-cause diagnosis + consistency checking ", "— the spec's Root Cause Score; penalizes a prerequisite edge's confidence when a strong topic has an evidenced-weak prerequisite (teacher-mediated).")
bullet("Misconception modeling ", "— LLM proposes a structured candidate on repeated misses; teacher confirms; verification-item assignment nudges confidence.")
bullet("Next-best-action ranking ", "— 6-factor composite + distinct action types (practice / diagnostic / concept / alternative / prereq-review / spaced-review / escalation).")
bullet("Governance ", "— nightly refit writes candidate snapshots (never auto-activates); ministry_admin review/promote/reject queue with rollback.")
bullet("Evaluation harness ", "— prequential forward-chaining AUC / log-loss / Brier / ECE; reports insufficient_data honestly.")
bullet("Explainability ", "— 5-part explanation endpoint with server-side IDOR-safe ownership check; state-snapshot history.")
bullet("Blocked-learning recovery ", "— detects 3 consecutive misses, routes through similar_to/alternative_to edges, escalates to teacher after 2 failed routes.")
doc.add_page_break()

# ===================== 5. PRODUCTION ROADMAP =====================
h1("5.  Production Completion Roadmap (Phases 0-12)")
para("Landed on this branch. Backend: go build & go vet clean. Frontend: type-check & lint clean. ai-service: imports clean. Not yet exercised against a live Docker stack.", bold=True, color=GREEN)
table(
    ["Phase", "Deliverable", "Status"],
    [
        ["1  Curriculum", "ApproveAndPromote batch refactor (pgx.Batch + Neo4j UNWIND); ~200s -> a handful of round-trips", "Built"],
        ["1  Curriculum", "S3StorageProvider + storage factory (STORAGE_PROVIDER env)", "Built"],
        ["1  Curriculum", "POST /curriculum/clos/resync — full CLO Neo4j resync", "Built"],
        ["2  Exam", "Lifecycle close endpoint + autosave draft + draft restore (V051/V052)", "Built"],
        ["3  EG-GCKT", "E2E verification against live stack", "Pending"],
        ["4  Student UX", "ExamAvailability / PreExam / SkillState pages, autosave + insight polling", "Built"],
        ["5  Dashboards", "HeatmapGrid -> ExplainPage click-through", "Built"],
        ["6  Security", "SecurityHeaders + RequestID, rate limiter, audit log (V053), RLS (V054)", "Built"],
        ["7  Observability", "pkg/metrics /metrics + /health on :9090, queue-depth poller, OTel telemetry", "Built"],
        ["10 Governance", "Model-governance review UI + snapshot lifecycle", "Built"],
        ["11 School Box", "sync-agent Go binary (outbox push / cloud pull, health :8081)", "Built"],
        ["12 Reporting", "reports domain + report_worker (V055); dead-letter handling (V056)", "Built"],
    ],
    widths=[1.3, 4.2, 0.9],
    status_col=2,
)
para("Resilience additions", bold=True, color=BLUE)
bullet("Dead-letter jobs ", "— failed jobs across curriculum/gap/kt workers land in public.dead_letter_jobs; nightly requeue in refit_worker.")
bullet("Real embedding provider ", "— OllamaEmbeddingProvider (nomic-embed-text, 768-dim) selectable via EMBEDDING_PROVIDER=stub|ollama|gemini, falls back to stub on connection error.")
doc.add_page_break()

# ===================== 6. KNOWN GAPS =====================
h1("6.  Known Gaps & Remaining Work")
table(
    ["Item", "State", "What's needed"],
    [
        ["EG-GCKT live verification", "Unverified", "Run against docker compose + real browser; nothing Playwright-tested yet"],
        ["Production roadmap live verification", "Unverified", "Compiles/type-checks; not yet smoke-tested end-to-end on a running stack"],
        ["Embedding semantics", "Partial", "Ollama provider exists but nomic-embed-text not benchmarked on Amharic/curriculum text"],
        ["Career matching", "Check", "Roadmap wired the route; confirm career_matcher/service.py is fully implemented, not a stub"],
        ["DKT / MIRT calibration", "Deferred", "Deliberately not built — no real interaction data yet (spec's own Phase 6 criterion)"],
        ["RLS enforcement", "Partial", "Policies added (V054); Supabase service_role bypasses — verify app connects with the intended role"],
        ["Migrations applied", "Pending", "V051-V056 written; run Flyway against Supabase to apply"],
    ],
    widths=[1.9, 1.1, 3.4],
    status_col=1,
)
para("Migrations must be applied", bold=True, color=AMBER)
para("V051-V056 are written but not yet run against Supabase. Apply via the standalone Flyway container (see Development Rules in CLAUDE.md) before the new endpoints will work against live data.")

# ===================== 7. IMMEDIATE NEXT STEPS =====================
h1("7.  Immediate Next Steps")
para("1. Apply migrations V051-V056 to Supabase via Flyway.", bold=True)
para("2. Bring up the full stack (docker compose up -d) and smoke-test the new surfaces: batch approval latency, exam close/autosave/draft, /metrics + /health, rate limiting, reports queue, sync-agent health.")
para("3. Playwright-verify EG-GCKT end-to-end (the largest unverified surface): trace on graded attempt, skill-state page, explain page, model-governance promote/reject.")
para("4. Confirm the Ollama embedding path end-to-end and decide stub vs. ollama vs. gemini for the default.")
para("5. Verify career_matcher service is a real implementation, not a stub.")

# ===================== 8. VERIFICATION STATUS =====================
h1("8.  Verification Status (this session)")
table(
    ["Check", "Command", "Result"],
    [
        ["Go build", "go build ./...", "Pass"],
        ["Go vet", "go vet ./...", "Pass"],
        ["Frontend types", "npm run type-check", "Pass"],
        ["Frontend lint", "npm run lint (--max-warnings 0)", "Pass"],
        ["ai-service import", "python -c 'import app.main'", "Pass (8 workers wired)"],
        ["Live stack smoke test", "docker compose up + browser", "Not yet run"],
    ],
    widths=[1.8, 2.6, 2.0],
    status_col=2,
)

para("End of report.", italic=True, color=GREY)

doc.save(r"D:\EDUGRAPH PROJECT\edugraph\progress.docx")
print("progress.docx regenerated OK")


